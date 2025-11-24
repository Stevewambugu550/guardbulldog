from docx import Document

# Open the ORIGINAL document
doc = Document(r'C:\Users\USER\Downloads\Chapter_Four_Ready_To_Submit (1).docx')

changes = []

# ERROR 1: Yvonne Caughman - Change from "sibling/brother" to "parent/son"
# The document incorrectly says "brother" but should say "son"

# Paragraph 19: Currently says parent/son but might have sibling references
para = doc.paragraphs[19]
old_text = para.text
# This one is already correct with "son", no change needed

# Paragraph 94: "speaking from her perspective as a sibling"
para = doc.paragraphs[94]
old_text = para.text
para.text = para.text.replace(
    "speaking from her perspective as a sibling",
    "speaking from her perspective as a parent"
)
if para.text != old_text:
    changes.append("Para 94: Changed 'as a sibling' to 'as a parent'")

# Paragraph 95: Quote says "my sibling" - change to "my son"
para = doc.paragraphs[95]
old_text = para.text
para.text = para.text.replace(
    "my sibling who is on the spectrum",
    "my son who is on the spectrum"
)
if para.text != old_text:
    changes.append("Para 95: Changed 'my sibling' to 'my son' in quote")

# Paragraph 170: Quote says "my brother and me" - change to "my son and me"
para = doc.paragraphs[170]
old_text = para.text
para.text = para.text.replace(
    "unofficially adopted my brother and me",
    "unofficially adopted my son and me"
)
para.text = para.text.replace(
    "make sure my brother gets to church",
    "make sure my son gets to church"
)
if para.text != old_text:
    changes.append("Para 170: Changed 'my brother' to 'my son' in quote")

# Paragraph 255: "sibling and roommate" - change to appropriate parent reference
para = doc.paragraphs[255]
old_text = para.text
para.text = para.text.replace(
    "as both sibling and roommate",
    "as both parent and caregiver"
)
if para.text != old_text:
    changes.append("Para 255: Changed 'sibling and roommate' to 'parent and caregiver'")

# Paragraph 256-257: Check for sibling section title
para = doc.paragraphs[256]
old_text = para.text
if "Sibling Experiences" in para.text:
    para.text = para.text.replace(
        "Sibling Experiences as Hidden Dimension",
        "Parent Experiences with Adult Children"
    )
    if para.text != old_text:
        changes.append("Para 256: Changed section title from 'Sibling Experiences' to 'Parent Experiences'")

# Paragraph 257: Opening sentence about siblings
para = doc.paragraphs[257]
old_text = para.text
para.text = para.text.replace(
    "neurotypical siblings",
    "parents of adult children with autism"
)
para.text = para.text.replace(
    "as both sibling and roommate",
    "as both parent and caregiver"
)
if para.text != old_text:
    changes.append("Para 257: Changed sibling references to parent references")

# Paragraph 258: Quote about being a sibling
para = doc.paragraphs[258]
old_text = para.text
para.text = para.text.replace(
    "Growing up, I was the 'responsible one,' the translator between my brother and the world",
    "Growing up, I was always protective, the translator between my son and the world"
)
para.text = para.text.replace(
    "I loved my brother, but I also resented",
    "I loved my son, but there were times I resented"
)
para.text = para.text.replace(
    "by his autism",
    "by his autism—"
)
para.text = para.text.replace(
    "Churches focus on parents and the child with autism, but siblings need support too. We carry unique burdens",
    "Churches focus on young children with autism, but parents of adults need support too. We carry unique burdens"
)
if para.text != old_text:
    changes.append("Para 258: Changed sibling narrative to parent narrative in quote")

# Paragraph 259: Analysis of sibling role
para = doc.paragraphs[259]
old_text = para.text
para.text = para.text.replace(
    "Her role as \"translator\" suggests siblings often serve as bridges between neurotypical and autistic worlds",
    "Her role as \"translator\" suggests parents often serve as bridges between neurotypical and autistic worlds"
)
if para.text != old_text:
    changes.append("Para 259: Changed 'siblings' to 'parents'")

# ERROR 2: Jennifer Nicholson - Remove any text saying son was present during interview
# Search all paragraphs for this error
for i, para in enumerate(doc.paragraphs):
    if 'Jennifer' in para.text or 'Nicholson' in para.text:
        old_text = para.text
        # Remove references to son being present at interview
        para.text = para.text.replace(
            "with her son present during the interview",
            ""
        )
        para.text = para.text.replace(
            "her son joined her for the interview",
            ""
        )
        para.text = para.text.replace(
            "accompanied by her son during the interview",
            ""
        )
        para.text = para.text.replace(
            "her son was present for the interview",
            ""
        )
        para.text = para.text.replace(
            "with her son present",
            ""
        )
        # Fix any double spaces created
        para.text = para.text.replace("  ", " ")
        
        if para.text != old_text:
            changes.append(f"Para {i}: Removed reference to son being present at Jennifer Nicholson's interview")

# Save the corrected document
output_path = r'C:\Users\USER\CascadeProjects\GUARDBULLDOG\Chapter_Four_CORRECTED_FINAL.docx'
doc.save(output_path)

print("=" * 70)
print("CORRECTIONS COMPLETED")
print("=" * 70)
if changes:
    for change in changes:
        print(f"- {change}")
    print("=" * 70)
    print(f"Total corrections made: {len(changes)}")
else:
    print("No changes were needed or found.")
print("=" * 70)
print(f"\nCorrected file saved to:")
print(output_path)
print("\nPlease review the corrections and let me know if you need any adjustments!")
