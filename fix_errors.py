from docx import Document

# Open the document
doc = Document(r'C:\Users\USER\Downloads\Chapter_Four_Ready_To_Submit (1).docx')

changes = []

# Paragraph 19: Change from parent/mother/son to sibling/sister/brother
para = doc.paragraphs[19]
old_text = para.text
para.text = para.text.replace(
    "Yvonne Caughman provided a unique parent perspective, sharing her experience as mother and caregiver of her son with autism in New York City.",
    "Yvonne Caughman provided a unique sibling perspective, sharing her experience as sister and roommate of her brother with autism in New York City."
)
para.text = para.text.replace(
    "how parents often serve as informal advocates and support providers within faith communities",
    "how siblings often serve as informal advocates and support providers within faith communities"
)
para.text = para.text.replace(
    "navigating city churches with her son",
    "navigating city churches with her brother"
)
if para.text != old_text:
    changes.append("Paragraph 19: Fixed Yvonne Caughman description from 'parent/mother/son' to 'sibling/sister/brother'")

# Paragraph 96: Change "her son's spiritual gifts" to "her brother's spiritual gifts"
para = doc.paragraphs[96]
old_text = para.text
para.text = para.text.replace(
    "Ms. Caughman's observation about her son's spiritual gifts",
    "Ms. Caughman's observation about her brother's spiritual gifts"
)
if para.text != old_text:
    changes.append("Paragraph 96: Fixed 'her son's spiritual gifts' to 'her brother's spiritual gifts'")

# Paragraph 122: Check the quote text
para = doc.paragraphs[122]
if "Yvonne Caughman shared how her family collaborated with church leadership around her son's needs:" in para.text:
    old_text = para.text
    para.text = para.text.replace(
        "around her son's needs:",
        "around her brother's needs:"
    )
    if para.text != old_text:
        changes.append("Paragraph 122: Fixed 'her son's needs' to 'her brother's needs'")

# Paragraph 124: Change "her son's need" to "her brother's need"
para = doc.paragraphs[124]
old_text = para.text
para.text = para.text.replace(
    "The accommodation of her son's need for routine",
    "The accommodation of her brother's need for routine"
)
if para.text != old_text:
    changes.append("Paragraph 124: Fixed 'her son's need' to 'her brother's need'")

# Save the corrected document
output_path = r'C:\Users\USER\CascadeProjects\GUARDBULLDOG\Chapter_Four_CORRECTED.docx'
doc.save(output_path)

print("=" * 70)
print("CORRECTIONS COMPLETED")
print("=" * 70)
if changes:
    for change in changes:
        print(f"- {change}")
    print("=" * 70)
    print(f"Total corrections made: {len(changes)}")
else:
    print("No changes were needed or found.")
print("=" * 70)
print(f"\nCorrected file saved to:\n{output_path}")
print("\nYou can now copy this file back to your Downloads folder.")
